import sys
import os
import json
import pytesseract
from PIL import Image
from docx import Document

# Set Tesseract configuration
# Assumes TESSDATA_PREFIX is set or tessdata is in current directory
if os.path.exists('tessdata'):
    os.environ['TESSDATA_PREFIX'] = os.path.abspath('tessdata')

def process_document(input_path, output_path):
    try:
        # Load image
        image = Image.open(input_path)
        
        # Perform OCR
        # lang='tam' for Tamil
        text = pytesseract.image_to_string(image, lang='tam')
        
        # Create Word Document
        doc = Document()
        doc.add_heading('Converted Tamil Document', 0)
        doc.add_paragraph(text)
        
        # Save Document
        doc.save(output_path)
        
        # Output result as JSON
        result = {
            "status": "completed",
            "extractedText": text,
            "convertedPath": output_path
        }
        print(json.dumps(result))
        
    except Exception as e:
        error_result = {
            "status": "failed",
            "error": str(e)
        }
        print(json.dumps(error_result))
        sys.exit(1)

if __name__ == "__main__":
    if len(sys.argv) < 3:
        print(json.dumps({"status": "failed", "error": "Missing arguments"}))
        sys.exit(1)
        
    input_file = sys.argv[1]
    output_file = sys.argv[2]
    
    process_document(input_file, output_file)
